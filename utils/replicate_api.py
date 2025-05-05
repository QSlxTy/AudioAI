import asyncio
import os
import random
from datetime import datetime

import replicate
from docx import Document

from bot_start import logger
from integrations.database.models.decoding import create_decoding_db, update_decoding_db
from src.config import Configuration

REPLICATE_API_TOKEN = Configuration.replicate_token

os.environ["REPLICATE_API_TOKEN"] = REPLICATE_API_TOKEN


async def replicate_func_url(audio, lang, speakers, words, format_decoding, user_id, session_maker):
    logger.info(f"Start decoding")
    replicate_id = int(random.randint(1, 999999))
    logger.info(f"Replicate settings: {user_id} - {audio} - {lang} - {speakers} - {words} - {format_decoding}")
    if lang == '':
        if speakers == '':
            output = await replicate.predictions.async_create(
                "cbd15da9f839c5f932742f86ce7def3a03c22e2b4171d42823e83e314547003f",
                input={
                    "file": audio,
                    "prompt": words,
                    "group_segments": True,
                    "offset_seconds": 0,
                    "transcript_output_format": "segments_only"
                },

            )
        else:
            output = await replicate.predictions.async_create(
                "cbd15da9f839c5f932742f86ce7def3a03c22e2b4171d42823e83e314547003f",
                input={
                    "file": audio,
                    "prompt": words,
                    "group_segments": True,
                    "num_speakers": int(speakers),
                    "offset_seconds": 0,
                    "transcript_output_format": "segments_only"
                }
            )
    else:
        if speakers == '':
            output = await replicate.predictions.async_create(
                "cbd15da9f839c5f932742f86ce7def3a03c22e2b4171d42823e83e314547003f",
                input={
                    "file": audio,
                    "prompt": words,
                    "language": lang,
                    "group_segments": True,
                    "offset_seconds": 0,
                    "transcript_output_format": "segments_only"
                }
            )
        else:
            output = await replicate.predictions.async_create(
                "cbd15da9f839c5f932742f86ce7def3a03c22e2b4171d42823e83e314547003f",
                input={
                    "file": audio,
                    "prompt": words,
                    "language": lang,
                    "group_segments": True,
                    "num_speakers": int(speakers),
                    "offset_seconds": 0,
                    "transcript_output_format": "segments_only"
                }
            )
    await create_decoding_db(user_id, 'in_process', output.id, session_maker)
    while True:
        try:
            await output.async_reload()
        except Exception:
            logger.warning(f'Get Prediction error --> {user_id} | {output.id} | {output.status}')
            await asyncio.sleep(90)
            raise
        if output.status == 'succeeded':
            await update_decoding_db(user_id, output.id, {'status': 'done'}, session_maker)
            logger.info(f"Decoding success | {output}")
            doc = Document()
            if 'Расшифровка с тайм-кодами и разбиением на спикеров' == format_decoding:
                def format_time(seconds):
                    hours = int(seconds // 3600)
                    minutes = int((seconds % 3600) // 60)
                    seconds = seconds % 60
                    return f"{hours:02}:{minutes:02}:{seconds:05.2f}".replace('.', ':')

                output = output.output
                logger.info(output)
                for segment in output['segments']:
                    start_time = format_time(segment['start'])
                    end_time = format_time(segment['end'])
                    speaker = segment['speaker']
                    text = segment['text']
                    doc.add_paragraph(f"[{start_time} - {end_time} - {speaker}]")
                    doc.add_paragraph(text)
                    doc.add_paragraph("")
            doc.save(f"media/{user_id}/{str(datetime.now()).split(' ')[0]}_СлушАЙ.docx")
            return f"media/{user_id}/{str(datetime.now()).split(' ')[0]}_СлушАЙ.docx"
        elif output.status != 'failed' and output.status != 'succeeded':
            logger.info(f'WAIT --> {user_id} | {output.id} | {output.status}')
            await asyncio.sleep(90)
        else:
            await update_decoding_db(user_id, output.id, {'status': 'failed'}, session_maker)
            logger.error(f'{user_id} | {output.id} | {output.status}')
            return None
