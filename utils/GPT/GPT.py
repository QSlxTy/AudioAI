from datetime import datetime

from docx import Document
from openai import OpenAI

from src.config import BotConfig
from utils.GPT.prompts import prompt_meeting_minutes, prompt_analysis, prompt_summary, prompt_tasks


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


client = OpenAI(
    api_key=BotConfig.gpt_token)


async def create_theses_from_transcription(user_id: int, summary_list: list, decode_path):
    list_doc = []
    for index,summary_format in enumerate(summary_list):
        text = extract_text_from_docx(decode_path)
        if summary_format == 'Протокол встречи':
            summary_format = prompt_meeting_minutes
        elif summary_format == 'Саммари':
            summary_format = prompt_summary
        elif summary_format == 'Перечень действий и задач':
            summary_format = prompt_analysis
        else:
            summary_format = prompt_tasks
        prompt = f"{summary_format}. Вот представленная расшифровка: {text}"

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ]
        )
        doc = Document()
        doc.add_paragraph(response.choices[0].message.content)
        doc.save(f"media/{user_id}/{str(datetime.now()).split(' ')[0]}_{summary_list[index]}_СлушАЙ.docx")
        list_doc.append(f"media/{user_id}/{str(datetime.now()).split(' ')[0]}_{summary_list[index]}_СлушАЙ.docx")
    return list_doc
